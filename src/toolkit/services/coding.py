"""Coding & thematic analysis — deductive, inductive, and hybrid."""

import asyncio
import json
import logging
import re
from collections import Counter
from collections.abc import AsyncIterator
from pathlib import Path

import pandas as pd

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _codebook_prompt_block(codebook: dict) -> str:
    lines = []
    for label, code in codebook.items():
        inc = "; ".join(code.get("inclusion_criteria", []))
        exc = "; ".join(code.get("exclusion_criteria", []))
        defn = code.get("definition", "")
        lines.append(f"CODE: {label.upper()}\n  Definition: {defn}\n  Include when: {inc}\n  Exclude when: {exc}")
    return "\n\n".join(lines)


async def _llm_text(client, model: str, system: str, user: str, max_tokens: int = 300) -> str:
    response = await client.chat.completions.create(
        model=model,
        messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
        temperature=0.1,
        max_tokens=max_tokens,
    )
    return (response.choices[0].message.content or "").strip()


# ---------------------------------------------------------------------------
# Deductive coding
# ---------------------------------------------------------------------------


async def run_deductive(
    chunks: list[dict],
    codebook: dict,
    client,
    model: str,
) -> AsyncIterator[str]:
    system = (
        "You are a qualitative research assistant specializing in deductive coding. "
        "Your task is to analyze text segments and identify which codes from the codebook apply.\n\n"
        + _codebook_prompt_block(codebook)
        + "\n\nCODING INSTRUCTIONS:\n"
        "1. Read each text segment carefully\n"
        "2. Apply ALL relevant codes from the codebook\n"
        "3. Only use codes that are explicitly defined above\n"
        "4. Return codes as a comma-separated list (e.g., 'CODE1,CODE2,CODE3')\n"
        "5. If no codes apply, return 'NO_CODES'\n"
        "6. Be consistent — similar content should receive similar codes\n"
        "7. Focus on manifest content (what is explicitly stated)\n\n"
        "Return only the comma-separated codes, no explanation needed."
    )
    valid_labels = set(codebook.keys())
    total = len(chunks)
    results = []

    for i, chunk in enumerate(chunks):
        if (i + 1) % 10 == 0 or i == 0:
            yield f"data: {json.dumps({'type': 'progress', 'message': f'Deductive coding chunk {i + 1}/{total}...'})}\n\n"

        try:
            raw = await _llm_text(client, model, system, chunk["text"], max_tokens=150)
            if raw.upper() == "NO_CODES":
                codes = []
            else:
                codes = [c.strip().lower() for c in raw.split(",") if c.strip().lower() in valid_labels]
        except Exception as e:
            logger.warning("Deductive coding failed for chunk %d: %s", i, e)
            codes = []

        results.append({"chunk_id": chunk.get("chunk_id", i + 1), "deductive_codes": codes})
        await asyncio.sleep(0)

    yield f"data: {json.dumps({'type': 'deductive_done', 'results': results})}\n\n"


# ---------------------------------------------------------------------------
# Inductive coding
# ---------------------------------------------------------------------------


def _filter_novel_codes(inductive_codes: dict, existing_labels: list[str], threshold: float = 0.75) -> dict:
    """Remove inductive codes that overlap too much with existing deductive codes."""
    try:
        from sentence_transformers import SentenceTransformer
        from sklearn.metrics.pairwise import cosine_similarity

        st_model = SentenceTransformer("all-MiniLM-L6-v2")
        ind_labels = list(inductive_codes.keys())
        ind_texts = [f"{k} {v.get('definition', '')}" for k, v in inductive_codes.items()]
        ded_texts = existing_labels  # label strings are enough

        ind_emb = st_model.encode(ind_texts, show_progress_bar=False)
        ded_emb = st_model.encode(ded_texts, show_progress_bar=False)
        sims = cosine_similarity(ind_emb, ded_emb)  # (n_ind, n_ded)

        novel = {}
        for i, label in enumerate(ind_labels):
            if sims[i].max() < threshold:
                novel[label] = inductive_codes[label]
            else:
                logger.info(
                    "Filtered inductive code '%s' — too similar to deductive code '%s'",
                    label,
                    existing_labels[int(sims[i].argmax())],
                )
        return novel
    except Exception as e:
        logger.warning("Novelty filter failed: %s — keeping all inductive codes", e)
        return inductive_codes


async def run_inductive(
    chunks: list[dict],
    existing_codes: list[str],
    client,
    model: str,
) -> AsyncIterator[str]:
    yield f"data: {json.dumps({'type': 'progress', 'message': 'Discovering inductive codes from sample...'})}\n\n"
    await asyncio.sleep(0)

    sample = chunks[:50]
    sample_text = "\n\n".join(f"[Chunk {c.get('chunk_id', i + 1)}]: {c['text']}" for i, c in enumerate(sample))
    existing_str = ", ".join(existing_codes) if existing_codes else "none"

    discovery_prompt = (
        "You are conducting INDUCTIVE CODING on interview transcripts.\n"
        "Your task is to identify EMERGENT THEMES that are NOT captured by the existing deductive codes.\n\n"
        f"EXISTING DEDUCTIVE CODES: {existing_str}\n\n"
        "SAMPLE CHUNKS FOR ANALYSIS:\n"
        f"{sample_text}\n\n"
        "TASK: Identify 8-12 EMERGENT INDUCTIVE CODES that capture important patterns NOT covered by deductive codes.\n\n"
        "For each code provide:\n"
        "**INDUCTIVE CODE: [SHORT_NAME]**\n"
        "Definition: [Clear description]\n"
        "Rationale: [Why this is important]\n"
        'Example: "[Direct quote]"\n'
        "When to Apply: [Clear criteria]\n\n"
        "Ensure there is a blank line between codes."
    )

    try:
        response = await client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": discovery_prompt}],
            temperature=0.4,
            max_tokens=3000,
        )
        raw = response.choices[0].message.content or ""
        inductive_codes = _parse_inductive_codes(raw)
    except Exception as e:
        logger.warning("Inductive code discovery failed: %s", e)
        inductive_codes = {}

    # Validate novelty: drop inductive codes too similar to existing deductive codes
    if inductive_codes and existing_codes:
        inductive_codes = _filter_novel_codes(inductive_codes, existing_codes)
        yield f"data: {json.dumps({'type': 'progress', 'message': f'{len(inductive_codes)} novel inductive codes after deduplication...'})}\n\n"
        await asyncio.sleep(0)

    yield f"data: {json.dumps({'type': 'progress', 'message': f'Applying {len(inductive_codes)} inductive codes...'})}\n\n"
    await asyncio.sleep(0)

    # Apply inductive codes to all chunks
    if inductive_codes:
        code_block = "\n".join(
            f"- {name}: {info['definition']} | Apply when: {info['when_to_apply']}"
            for name, info in inductive_codes.items()
        )
        ind_system = (
            "Apply these INDUCTIVE CODES to text chunks.\n\n"
            f"INDUCTIVE CODES:\n{code_block}\n\n"
            "Instructions:\n"
            "1. Apply ONLY codes that clearly match\n"
            "2. Return codes as comma-separated list\n"
            "3. If no codes apply, return 'NONE'\n\n"
            "Return ONLY the code names."
        )
        total = len(chunks)
        ind_results = []
        for i, chunk in enumerate(chunks):
            if (i + 1) % 10 == 0 or i == 0:
                yield f"data: {json.dumps({'type': 'progress', 'message': f'Inductive coding chunk {i + 1}/{total}...'})}\n\n"
            try:
                raw = await _llm_text(client, model, ind_system, chunk["text"], max_tokens=100)
                if raw.upper() == "NONE":
                    codes = []
                else:
                    codes = [c.strip().lower() + "_ind" for c in raw.split(",") if c.strip()]
            except Exception as e:
                logger.warning("Inductive application failed for chunk %d: %s", i, e)
                codes = []
            ind_results.append({"chunk_id": chunk.get("chunk_id", i + 1), "inductive_codes": codes})
            await asyncio.sleep(0)
    else:
        ind_results = [{"chunk_id": c.get("chunk_id", i + 1), "inductive_codes": []} for i, c in enumerate(chunks)]

    yield f"data: {json.dumps({'type': 'inductive_done', 'results': ind_results, 'discovered_codes': inductive_codes})}\n\n"


def _parse_inductive_codes(text: str) -> dict:
    codes = {}
    blocks = re.split(r"\*\*INDUCTIVE CODE:", text)
    for block in blocks[1:]:
        name_match = re.match(r"\s*([^\*\n]+)", block)
        if not name_match:
            continue
        name = re.sub(r"[^a-zA-Z0-9_]", "_", name_match.group(1).strip())[:25].lower()
        defn_match = re.search(r"Definition:\s*(.+?)(?=\n\w|Rationale:|$)", block, re.DOTALL)
        when_match = re.search(r"When to Apply:\s*(.+?)(?=\n\n|\Z)", block, re.DOTALL)
        example_match = re.search(r'Example:\s*["\']?(.+?)["\']?\s*(?=When to Apply:|$)', block, re.DOTALL)
        codes[name] = {
            "definition": defn_match.group(1).strip() if defn_match else "",
            "when_to_apply": when_match.group(1).strip() if when_match else "",
            "example": example_match.group(1).strip() if example_match else "",
        }
    return codes


# ---------------------------------------------------------------------------
# Theme building
# ---------------------------------------------------------------------------


async def build_themes(
    df: pd.DataFrame,
    client,
    model: str,
) -> AsyncIterator[str]:
    yield f"data: {json.dumps({'type': 'progress', 'message': 'Analysing code patterns...'})}\n\n"
    await asyncio.sleep(0)

    all_codes: list[str] = []
    deductive_codes: list[str] = []
    inductive_codes: list[str] = []

    for _, row in df.iterrows():
        ded = row.get("deductive_codes", []) or []
        ind = row.get("inductive_codes", []) or []
        if isinstance(ded, str):
            ded = [c for c in ded.split(",") if c.strip()]
        if isinstance(ind, str):
            ind = [c for c in ind.split(",") if c.strip()]
        all_codes.extend(ded + ind)
        deductive_codes.extend(ded)
        inductive_codes.extend(ind)

    top_ded = Counter(deductive_codes).most_common(10)
    top_ind = Counter(inductive_codes).most_common(10)
    top_all = Counter(all_codes).most_common(20)

    # Sample coded chunks for context
    sample_chunks = df.head(3)[["chunk_id", "text", "deductive_codes", "inductive_codes"]].to_dict("records")
    sample_text = "\n\n".join(
        f"Chunk {r['chunk_id']}: {r['text'][:300]}...\n  Deductive: {r['deductive_codes']}\n  Inductive: {r['inductive_codes']}"
        for r in sample_chunks
    )

    top_ded_str = "\n".join(f"  {code}: {freq}" for code, freq in top_ded)
    top_ind_str = "\n".join(f"  {code}: {freq}" for code, freq in top_ind)

    theme_prompt = (
        "Du er en ekspert i kvalitativ forskning og skal opbygge TEMAER ud fra mixed-method kodningsresultater.\n\n"
        f"KODNINGSOVERSIGT:\n"
        f"- Antal kodeappliceringer i alt: {len(all_codes)}\n"
        f"- Unikke koder: {len(set(all_codes))}\n\n"
        f"TOP DEDUKTIVE KODER:\n{top_ded_str}\n\n"
        f"TOP INDUKTIVE KODER:\n{top_ind_str}\n\n"
        f"EKSEMPEL PÅ KODEDE UDDRAG:\n{sample_text}\n\n"
        "OPGAVE: Udarbejd 5-7 HIERARKISKE TEMAER, som:\n"
        "1. Integrerer indsigter fra både deduktive og induktive koder\n"
        "2. Har klare hovedtemaer med 2-3 undertemaer hver\n"
        "3. Er handlingsorienterede og relevante\n\n"
        "Skriv svaret på DANSK. Formater hvert tema som:\n\n"
        "TEMA [Nummer]: [Klart og beskrivende navn]\n"
        "Kernekonceptet: [2-3 sætninger der forklarer, hvad temaet indfanger]\n"
        "Undertemaer:\n"
        "  a) [Undertema navn]: [Kort beskrivelse]\n"
        "  b) [Undertema navn]: [Kort beskrivelse]\n"
        "Nøglefund: [Den vigtigste indsigt dette tema afslører]\n"
        "Evidensstyrke: [Stærk/Moderat/Fremvoksende - baseret på kodehyppighed]"
    )

    yield f"data: {json.dumps({'type': 'progress', 'message': 'Generating thematic analysis...'})}\n\n"
    await asyncio.sleep(0)

    try:
        response = await client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": theme_prompt}],
            temperature=0.3,
            max_tokens=4000,
        )
        themes_text = response.choices[0].message.content or ""
    except Exception as e:
        logger.error("Theme building failed: %s", e)
        themes_text = f"Theme generation failed: {e}"

    yield f"data: {json.dumps({'type': 'done', 'themes': themes_text, 'code_frequencies': dict(top_all)})}\n\n"


# ---------------------------------------------------------------------------
# Export helpers
# ---------------------------------------------------------------------------


def merge_coding_results(
    chunks: list[dict],
    deductive_results: list[dict],
    inductive_results: list[dict],
) -> pd.DataFrame:
    import datetime

    ded_map = {r["chunk_id"]: r["deductive_codes"] for r in deductive_results}
    ind_map = {r["chunk_id"]: r["inductive_codes"] for r in inductive_results}
    timestamp = datetime.datetime.now().isoformat()

    rows = []
    for c in chunks:
        cid = c.get("chunk_id", 0)
        ded = ded_map.get(cid, [])
        ind = ind_map.get(cid, [])
        all_codes = ded + ind
        text = c["text"]
        words = text.split()
        word_count = len(words)
        try:
            import nltk

            sentence_count = len(nltk.sent_tokenize(text)) if text else 0
        except Exception:
            sentence_count = text.count(".") + text.count("?") + text.count("!") or 1
        rows.append(
            {
                "chunk_id": cid,
                "text": text,
                "word_count": word_count,
                "char_count": len(text),
                "sentence_count": sentence_count,
                "speaker": c.get("speaker", ""),
                "avg_words_per_sentence": round(word_count / sentence_count, 1) if sentence_count else 0,
                "source_file": c.get("source_file", ""),
                "processed_timestamp": timestamp,
                "deductive_codes": ", ".join(ded),
                "inductive_codes": ", ".join(ind),
                "all_codes": ", ".join(all_codes),
                "total_code_count": len(all_codes),
                "coding_status": _coding_status(ded, ind),
            }
        )
    return pd.DataFrame(rows)


def _coding_status(ded: list, ind: list) -> str:
    if ded and ind:
        return "Both"
    if ded:
        return "Deductive"
    if ind:
        return "Inductive"
    return "No codes"


def save_coding_results(
    df: pd.DataFrame,
    session_path: Path,
    codebook: dict | None = None,
    inductive_codes: dict | None = None,
) -> Path:
    out = session_path / "coded_data.xlsx"

    # Code frequencies with type annotation
    ded_freq: Counter[str] = Counter()
    ind_freq: Counter[str] = Counter()
    for _, row in df.iterrows():
        for c in (row.get("deductive_codes") or "").split(","):
            c = c.strip()
            if c:
                ded_freq[c] += 1
        for c in (row.get("inductive_codes") or "").split(","):
            c = c.strip()
            if c:
                ind_freq[c] += 1

    freq_rows = [{"Code": k, "Type": "Deductive", "Frequency": v} for k, v in ded_freq.most_common()]
    freq_rows += [{"Code": k, "Type": "Inductive", "Frequency": v} for k, v in ind_freq.most_common()]
    freq_rows.sort(key=lambda x: x["Frequency"] if isinstance(x["Frequency"], int) else 0, reverse=True)
    freq_df = pd.DataFrame(freq_rows)

    # Code co-occurrence (combinations)
    combo_counter: Counter[str] = Counter()
    for _, row in df.iterrows():
        all_c = [c.strip() for c in (row.get("all_codes") or "").split(",") if c.strip()]
        if len(all_c) >= 2:
            for i in range(len(all_c)):
                for j in range(i + 1, len(all_c)):
                    pair = f"{all_c[i]} + {all_c[j]}"
                    combo_counter[pair] += 1
    combo_df = pd.DataFrame([{"Combination": k, "Frequency": v} for k, v in combo_counter.most_common(30)])

    # Deductive codebook sheet
    if codebook:
        cb_rows = []
        for label, code in codebook.items():
            inc = code.get("inclusion_criteria", [])
            exc = code.get("exclusion_criteria", [])
            examples = code.get("examples", [])
            cb_rows.append(
                {
                    "code_label": label,
                    "definition": code.get("definition", ""),
                    "inclusion_criteria": "; ".join(inc) if isinstance(inc, list) else inc,
                    "exclusion_criteria": "; ".join(exc) if isinstance(exc, list) else exc,
                    "example_1": examples[0] if len(examples) > 0 else "",
                    "example_2": examples[1] if len(examples) > 1 else "",
                    "frequency": ded_freq.get(label, 0),
                }
            )
        codebook_df = pd.DataFrame(cb_rows)
    else:
        codebook_df = None

    # Inductive codes sheet
    if inductive_codes:
        ind_rows = []
        for code_name, info in inductive_codes.items():
            ind_rows.append(
                {
                    "Code": code_name,
                    "Definition": info.get("definition", ""),
                    "Application": info.get("when_to_apply", ""),
                    "Example": info.get("example", ""),
                    "Frequency": ind_freq.get(code_name + "_ind", 0),
                }
            )
        ind_df = pd.DataFrame(ind_rows)
    else:
        ind_df = None

    summary = pd.DataFrame(
        {
            "Metric": ["Total chunks", "Coded chunks", "Deductive only", "Inductive only", "Both", "No codes"],
            "Count": [
                len(df),
                len(df[df["coding_status"] != "No codes"]),
                len(df[df["coding_status"] == "Deductive"]),
                len(df[df["coding_status"] == "Inductive"]),
                len(df[df["coding_status"] == "Both"]),
                len(df[df["coding_status"] == "No codes"]),
            ],
        }
    )

    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="Coded_Data", index=False)
        summary.to_excel(writer, sheet_name="Coding_Summary", index=False)
        if codebook_df is not None:
            codebook_df.to_excel(writer, sheet_name="Deductive_Codebook", index=False)
        if ind_df is not None:
            ind_df.to_excel(writer, sheet_name="Inductive_Codes", index=False)
        if not freq_df.empty:
            freq_df.to_excel(writer, sheet_name="Code_Frequencies", index=False)
        if not combo_df.empty:
            combo_df.to_excel(writer, sheet_name="Code_Combinations", index=False)
    return out


def save_themes_docx(themes_text: str, session_path: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_heading("Thematic Analysis", 0)

    for line in themes_text.splitlines():
        line = line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if re.match(r"^TEMA \d+:", line) or re.match(r"^THEME \d+:", line):
            doc.add_heading(line, level=1)
        elif re.match(r"^(Kernekonceptet|Nøglefund|Evidensstyrke|Core Concept|Key Finding|Evidence Strength):", line):
            p = doc.add_paragraph()
            p.add_run(line.split(":", 1)[0] + ":").bold = True
            p.add_run(line.split(":", 1)[1] if ":" in line else "")
        elif re.match(r"^\s+[a-z]\)", line):
            doc.add_paragraph(line.strip(), style="List Bullet")
        elif line in ("Undertemaer:", "Sub-themes:"):
            p = doc.add_paragraph()
            p.add_run(line).bold = True
        else:
            doc.add_paragraph(line)

    out = session_path / "themes_report.docx"
    doc.save(out)
    return out
